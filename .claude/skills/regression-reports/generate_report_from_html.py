#!/usr/bin/env python3
"""
Parse HTML regression tables and generate Word document.
"""

import re
from bs4 import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_variable_mapping(mapping_file):
    mapping = {}
    with open(mapping_file, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if ' → ' in line and not line.startswith('#'):
                parts = line.split(' → ')
                if len(parts) == 2:
                    mapping[parts[0].strip()] = parts[1].strip()
    return mapping


def apply_variable_mapping(text, mapping):
    for old_name, new_name in mapping.items():
        text = re.sub(r'\b' + re.escape(old_name) + r'\b', new_name, text)
    return text


def parse_html_table(html_file):
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    table = soup.find('table', class_='texreg')

    if not table:
        return None

    header_row = table.find('thead').find('tr')
    headers = [th.get_text(strip=True) for th in header_row.find_all('th')]
    dvs = [h for h in headers[1:] if h and h != ' ' and h != '  ' and h != '   ']

    tbody = table.find('tbody')
    rows = tbody.find_all('tr')

    coef_rows = []
    stat_rows = []
    in_stats = False

    i = 0
    while i < len(rows):
        row = rows[i]
        cells = row.find_all('td')

        if not cells:
            i += 1
            continue

        var_name = cells[0].get_text(strip=True)

        if 'Num. obs.' in var_name:
            in_stats = True

        if in_stats:
            values = [cell.get_text(strip=True) for cell in cells[1:]]
            stat_rows.append({
                'variable': var_name,
                'values': values
            })
            i += 1
        else:
            if var_name and var_name != '&nbsp;':
                coef_values = []
                for cell in cells[1:]:
                    text = cell.get_text(strip=True)
                    text = text.replace('&nbsp;', '')
                    coef_values.append(text)

                se_values = []
                if i + 1 < len(rows):
                    se_row = rows[i + 1]
                    se_cells = se_row.find_all('td')
                    for cell in se_cells[1:]:
                        text = cell.get_text(strip=True)
                        se_values.append(text)

                coef_rows.append({
                    'variable': var_name,
                    'coef_row': coef_values,
                    'se_row': se_values
                })
                i += 2
            else:
                i += 1

    return {
        'dvs': dvs,
        'coef_rows': coef_rows,
        'stat_rows': stat_rows
    }


def create_regression_table(doc, table_data, mapping, table_num, section_title, sample_desc, dv_type, iv_desc, is_first_in_section=True, column_group_dvs=None):
    if is_first_in_section:
        doc.add_heading(section_title, level=1)

    doc.add_paragraph(f"Table {table_num}")

    if not table_data['coef_rows']:
        return

    n_data_cols = len(table_data['coef_rows'][0]['coef_row'])
    n_cols = n_data_cols + 1
    n_rows = len(table_data['coef_rows']) * 2 + len(table_data['stat_rows']) + 2

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'

    header_row1 = table.rows[0].cells
    header_row1[0].text = ''

    dvs = column_group_dvs if column_group_dvs else table_data['dvs']
    cols_per_dv = n_data_cols // len(dvs)

    for i, dv in enumerate(dvs):
        mapped_dv = apply_variable_mapping(dv, mapping)
        start_col = i * cols_per_dv + 1
        end_col = start_col + cols_per_dv - 1

        if cols_per_dv > 1:
            merged_cell = header_row1[start_col].merge(header_row1[end_col])
        else:
            merged_cell = header_row1[start_col]
        merged_cell.text = f"DV: {mapped_dv}"
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    header_row2 = table.rows[1].cells
    header_row2[0].text = ''
    for i in range(n_data_cols):
        header_row2[i + 1].text = f"({i + 1})"
        header_row2[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_row = 2
    for row_data in table_data['coef_rows']:
        mapped_var = apply_variable_mapping(row_data['variable'], mapping)

        coef_cells = table.rows[current_row].cells
        coef_cells[0].text = mapped_var
        for col_idx, val in enumerate(row_data['coef_row']):
            if col_idx + 1 < n_cols:
                coef_cells[col_idx + 1].text = val
                coef_cells[col_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        se_cells = table.rows[current_row + 1].cells
        se_cells[0].text = ''
        for col_idx, val in enumerate(row_data['se_row']):
            if col_idx + 1 < n_cols:
                se_cells[col_idx + 1].text = val
                se_cells[col_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        current_row += 2

    for stat_data in table_data['stat_rows']:
        stat_cells = table.rows[current_row].cells
        stat_cells[0].text = stat_data['variable']
        for col_idx, val in enumerate(stat_data['values']):
            if col_idx + 1 < n_cols:
                stat_cells[col_idx + 1].text = val
                stat_cells[col_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        current_row += 1

    doc.add_paragraph()
    notes_para = doc.add_paragraph()
    notes_para.add_run("Notes:").bold = True
    doc.add_paragraph(f"Sample: {sample_desc}", style='List Bullet')
    doc.add_paragraph(f"DV: {dv_type}", style='List Bullet')
    doc.add_paragraph(f"IV: {iv_desc}", style='List Bullet')
    doc.add_paragraph()
